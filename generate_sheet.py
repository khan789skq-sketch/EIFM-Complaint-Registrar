from pathlib import Path
import shutil
import openpyxl
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).resolve().parent

def _first_existing(*paths):
    for p in paths:
        if p.exists():
            return p
    return paths[0]

EXCEL_3 = _first_existing(BASE_DIR / 'All-3.xlsx', BASE_DIR / 'templates' / 'All-3.xlsx')
EXCEL_1 = _first_existing(BASE_DIR / 'All.xlsx', BASE_DIR / 'templates' / 'All.xlsx')
WCC_TEMPLATE = _first_existing(
    BASE_DIR / '08 -EIFMEN08 - Work Completion Certificate.docx',
    BASE_DIR / 'wcc.docx',
    BASE_DIR / 'templates' / '08 -EIFMEN08 - Work Completion Certificate.docx',
)

def _sheet_names(path):
    if not path.exists(): return []
    wb = openpyxl.load_workbook(path, read_only=True, data_only=False)
    try: return list(wb.sheetnames)
    finally: wb.close()

EQUIPMENT_SHEETS=[]
for p in (EXCEL_3, EXCEL_1):
    for s in _sheet_names(p):
        if s not in EQUIPMENT_SHEETS: EQUIPMENT_SHEETS.append(s)

def _source_for_equipment(equipment):
    for p in (EXCEL_3, EXCEL_1):
        if p.exists() and equipment in _sheet_names(p): return p
    raise FileNotFoundError(f'Equipment sheet not found: {equipment}')

def read_template_tasks(equipment):
    src=_source_for_equipment(equipment)
    wb=openpyxl.load_workbook(src,data_only=False)
    try:
        ws=wb[equipment]; _,rows=_find_checklist_rows(ws)
        return [str(ws.cell(r,2).value) for r in rows if ws.cell(r,2).value]
    finally: wb.close()

def _find_checklist_rows(ws):
    header_row=None
    for r in range(1,ws.max_row+1):
        if str(ws.cell(r,1).value or '').strip().lower()=='sl. no.': header_row=r; break
    if header_row is None: return None,[]
    rows=[]
    for r in range(header_row+1,ws.max_row+1):
        a,b=ws.cell(r,1).value,ws.cell(r,2).value
        if isinstance(a,(int,float)) and b: rows.append(r)
        elif str(a or '').strip().upper().startswith('GENERAL NOTICE'): break
    return header_row,rows

def copy_selected_equipment_workbook(selected_equipment, output_path, source_paths=(EXCEL_3,EXCEL_1)):
    """Create ONE Excel file containing exactly the selected equipment sheets.
    Each sheet is copied from its original workbook; no checklist row limit is applied.
    """
    selected=list(dict.fromkeys(selected_equipment or []))
    if not selected: raise ValueError('Select at least one equipment.')
    output_path=Path(output_path); output_path.parent.mkdir(parents=True,exist_ok=True)
    temp=output_path.with_name(output_path.stem+'_source.xlsx')
    # Start with the first selected sheet's original workbook, then remove all unselected sheets.
    first_src=_source_for_equipment(selected[0])
    shutil.copy2(first_src,temp)
    wb=openpyxl.load_workbook(temp)
    try:
        # Copy sheets that are not in the starting workbook from their source workbooks.
        existing=set(wb.sheetnames)
        for eq in selected:
            src=_source_for_equipment(eq)
            if eq in existing: continue
            src_wb=openpyxl.load_workbook(src)
            try:
                src_ws=src_wb[eq]
                dst=wb.create_sheet(eq)
                for row in src_ws.iter_rows():
                    for cell in row:
                        nc=dst[cell.coordinate]
                        nc.value=cell.value
                        if cell.has_style:
                            from copy import copy
                            nc._style=copy(cell._style)
                        if cell.number_format: nc.number_format=cell.number_format
                        if cell.alignment: nc.alignment=copy(cell.alignment)
                        if cell.protection: nc.protection=copy(cell.protection)
                        if cell.hyperlink: nc._hyperlink=copy(cell.hyperlink)
                        if cell.comment: nc.comment=copy(cell.comment)
                for key,dim in src_ws.column_dimensions.items(): dst.column_dimensions[key]=copy(dim)
                for key,dim in src_ws.row_dimensions.items(): dst.row_dimensions[key]=copy(dim)
                for mr in src_ws.merged_cells.ranges: dst.merge_cells(str(mr))
                dst.sheet_view.showGridLines=src_ws.sheet_view.showGridLines
                existing.add(eq)
            finally: src_wb.close()
        for name in list(wb.sheetnames):
            if name not in selected: del wb[name]
        # Keep selected equipment order.
        for idx,name in enumerate(selected): wb._sheets.insert(idx,wb._sheets.pop(wb._sheets.index(wb[name])))
        wb.save(output_path)
    finally:
        wb.close()
        if temp.exists(): temp.unlink(missing_ok=True)
    return output_path

def _set_para(p,text,size=9):
    p.text=text
    for run in p.runs: run.font.size=Pt(size)

def _add_signature_image(paragraph,image_path,width=1.25):
    if not image_path: return
    run=paragraph.add_run(); run.add_picture(str(image_path),width=Inches(width))

def _insert_photo_row(doc,before_image,after_image):
    """Put both photos on ONE existing blank paragraph so they stay on one page."""
    if not (before_image or after_image): return
    target=None
    # Prefer the blank area immediately before Remarks / Suggestions.
    for p in doc.paragraphs:
        if not p.text.strip() and target is None: target=p
    # Better target: first blank paragraph after satisfaction text (usually index 31).
    for i,p in enumerate(doc.paragraphs):
        if 'Dear valued customer' in p.text and i+1 < len(doc.paragraphs):
            target=doc.paragraphs[i+1]; break
    if target is None: target=doc.paragraphs[-1]
    target.text=''
    target.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if before_image:
        r=target.add_run(); r.add_picture(str(before_image),width=Inches(2.55))
        r.add_text('   ')
    if after_image:
        r=target.add_run(); r.add_picture(str(after_image),width=Inches(2.55))

def create_ppm_from_template(
    equipment, project, location, unit, frequency, category, fiscal_year,
    wo_number, ppm_number, scheduled_month, service_date, time_start, time_finish,
    tasks, statuses, remarks, followups, technician, engineer, summary, output_path,
):
    """Backward-compatible single-equipment PPM generator."""
    return create_ppm_package([equipment],project,location,unit,frequency,category,fiscal_year,wo_number,ppm_number,scheduled_month,service_date,time_start,time_finish,tasks,statuses,remarks,followups,technician,engineer,summary,output_path)

def create_ppm_package(selected_equipment, project, location, unit, frequency, category, fiscal_year,
                       wo_number, ppm_number, scheduled_month, service_date, time_start, time_finish,
                       task_data, statuses, remarks, followups, technician, engineer, summary, output_path):
    selected=list(dict.fromkeys(selected_equipment or []))
    if not selected: raise ValueError('Select at least one equipment.')
    output_path=Path(output_path); output_path.parent.mkdir(parents=True,exist_ok=True)
    # Use first equipment's original workbook as base, then retain/copy all selected sheets.
    copy_selected_equipment_workbook(selected,output_path)
    wb=openpyxl.load_workbook(output_path)
    try:
        for equipment in selected:
            ws=wb[equipment]
            ws['C3']=project; ws['C4']=location; ws['C5']=unit; ws['C6']=frequency; ws['C7']=category; ws['C8']=equipment
            ws['K3']=fiscal_year; ws['K4']=wo_number; ws['K5']=scheduled_month; ws['K6']=service_date; ws['K7']=time_start; ws['K8']=time_finish
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value,str) and 'PPM Number' in cell.value:
                        ws.cell(cell.row,cell.column+1).value=ppm_number
            _,source_rows=_find_checklist_rows(ws)
            eq_tasks=(task_data or {}).get(equipment,[])
            for idx,r in enumerate(source_rows,start=1):
                ws.cell(r,7).value='✓' if (statuses or {}).get((equipment,idx)) else ''
                ws.cell(r,8).value='✓' if (statuses or {}).get((equipment,f'no_{idx}')) else ''
                ws.cell(r,9).value=(remarks or {}).get((equipment,idx),'')
                ws.cell(r,11).value=(followups or {}).get((equipment,idx),'')
                if idx<=len(eq_tasks): ws.cell(r,2).value=eq_tasks[idx-1]
            for r in range(1,ws.max_row+1):
                a=str(ws.cell(r,1).value or '')
                if a.startswith('Tech. Date/Sign:') and technician: ws.cell(r,1).value=f'Tech. Date/Sign: {technician}'
                if a.startswith('Eng./Sup Date Sign') and engineer: ws.cell(r,1).value=f'Eng./Sup Date Sign : {engineer}'
                if a.startswith('REPORT SUMMARY') and summary and r+1<=ws.max_row: ws.cell(r+1,1).value=summary
        wb.save(output_path)
    finally: wb.close()
    return output_path

def create_wcc_from_template(
    job_order, client, project, location, tel, details, completion,
    site_sig, site_name, site_date, site_id,
    hod_sig, hod_name, hod_date, hod_id,
    docs, client_sig, client_name, client_phone, satisfaction,
    remarks, client_sign_date, output_path,
    before_image=None, after_image=None, ppm_number=None, ppm_year=None,
):
    if not WCC_TEMPLATE.exists(): raise FileNotFoundError(f'WCC template not found: {WCC_TEMPLATE}')
    output_path=Path(output_path); output_path.parent.mkdir(parents=True,exist_ok=True)
    doc=Document(WCC_TEMPLATE)
    jo_text=f'{ppm_number} PPM Service Year {ppm_year or ""}'.strip() if ppm_number else (job_order or '')
    replacements={
        4:f'Job Order Number  :- {jo_text}',5:f'Client                         :- {client}',6:f'Project                       :- {project}',
        7:f'Location                     :- {location}                            Tel. No. :- {tel}',15:f'Date & Time of Completion :- {completion}',
        17:f'Signature of Site in Charge:-    Name :- {site_name}    Date :- {site_date}',18:f'  ID : - {site_id}',
        19:f'8.  Signature of HOD :-    Name:- {hod_name}    Date :- {hod_date}    ID :- {hod_id}',24:' Client Signature:-',26:f'Name: - {client_name}',28:f'Phone No. :- {client_phone}',39:f'Client Signature  ---------------------    Date :- {client_sign_date}',
    }
    for idx,value in replacements.items():
        if idx<len(doc.paragraphs): _set_para(doc.paragraphs[idx],value)
    lines=[x.strip() for x in (details or '').splitlines() if x.strip()]
    for idx in range(9,15):
        if idx<len(doc.paragraphs): _set_para(doc.paragraphs[idx],lines[idx-9] if idx-9<len(lines) else '')
    selected=set(docs or [])
    if len(doc.paragraphs)>21: _set_para(doc.paragraphs[21],'   '.join(f"{'☑' if x in selected else '☐'} {x}" for x in ['LPO','Invoice','Delivery Note','Petty Cash']))
    if len(doc.paragraphs)>23: _set_para(doc.paragraphs[23],'   '.join(f"{'☑' if x in selected else '☐'} {x}" for x in ['Material Requisition','Job Completion']))
    for p in doc.paragraphs:
        if 'Dear valued customer' in p.text: _set_para(p,p.text+f'   Selected: {satisfaction}')
        if p.text.strip()=='Remarks /Suggestions:': _set_para(p,p.text+f' {remarks or ""}')
    # Insert actual signature images at the three signature locations.
    if site_sig and len(doc.paragraphs)>17:
        p=doc.paragraphs[17]; p.text='Signature of Site in Charge:- '; _add_signature_image(p,site_sig)
        p.add_run(f'    Name :- {site_name}    Date :- {site_date}')
    if hod_sig and len(doc.paragraphs)>19:
        p=doc.paragraphs[19]; p.text='8.  Signature of HOD :- '; _add_signature_image(p,hod_sig)
        p.add_run(f'    Name:- {hod_name}    Date :- {hod_date}    ID :- {hod_id}')
    if client_sig and len(doc.paragraphs)>24:
        p=doc.paragraphs[24]; p.text=' Client Signature:- '; _add_signature_image(p,client_sig)
    if before_image or after_image: _insert_photo_row(doc,before_image,after_image)
    doc.save(output_path); return output_path

def generate_ppm(source,equipment,meta,out):
    return create_ppm_from_template(equipment,meta.get('project',''),meta.get('location',''),meta.get('unit',''),meta.get('frequency',''),meta.get('category',''),meta.get('fiscal_year',''),meta.get('wo',''),meta.get('ppm','1st PPM'),meta.get('month',''),meta.get('service_date',''),meta.get('time_start',''),meta.get('time_finish',''),meta.get('tasks',[]),meta.get('statuses',{}),meta.get('remarks',{}),meta.get('followups',{}),meta.get('technician',''),meta.get('engineer',''),meta.get('summary',''),out)

def generate_wcc(template,meta,out):
    return create_wcc_from_template(meta.get('job',''),meta.get('client',''),meta.get('project',''),meta.get('location',''),meta.get('tel',''),meta.get('details',''),meta.get('completion',''),meta.get('site_sig'),meta.get('site',''),meta.get('site_date',''),meta.get('siteid',''),meta.get('hod_sig'),meta.get('hod',''),meta.get('hod_date',''),meta.get('hodid',''),meta.get('docs',[]),meta.get('client_sig'),meta.get('cname',''),meta.get('phone',''),meta.get('satisfaction',''),meta.get('remarks',''),meta.get('client_sign_date',''),out)
